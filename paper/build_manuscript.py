# -*- coding: utf-8 -*-
"""
Assemble the revised manuscript as a Springer-style .docx for
Artificial Intelligence Review (SPBASIC author-year citations).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from references import REFERENCES, intext, sorted_reference_strings

C = {k: intext(k) for k in REFERENCES}  # short handle for in-text labels

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


def H1(text):
    h = doc.add_heading(text, level=1)
    return h


def H2(text):
    return doc.add_heading(text, level=2)


def H3(text):
    return doc.add_heading(text, level=3)


def P(text, justify=True, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def centered(text, bold=False, size=None, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def labelled(label, body):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(label + " ")
    r.bold = True
    p.add_run(body)
    return p


def make_table(headers, rows, caption=None, widths=None):
    if caption:
        cap = doc.add_paragraph()
        cr = cap.add_run(caption)
        cr.bold = True
        cr.font.size = Pt(10)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# =====================================================================
# TITLE BLOCK
# =====================================================================
title = doc.add_heading(
    "Breaking the Memory Wall: A Survey of Key-Value (KV) Cache Compression "
    "for Efficient Large Language Model (LLM) Inference", level=0)
for run in title.runs:
    run.font.size = Pt(16)

# Artificial Intelligence Review is single-blind: authors are named on the title page.
centered("Rohith Reddy B. C.", bold=True, size=12)
centered("[Department], [Institution], [City], [Country]", italic=True, size=10)
centered("Corresponding author: rohithreddybc98@gmail.com  ·  ORCID: 0000-0000-0000-0000", italic=True, size=10)
doc.add_paragraph()

# =====================================================================
# ABSTRACT
# =====================================================================
H1("Abstract")
P("The deployment of large language models (LLMs) at long context lengths is "
  "increasingly limited by memory rather than compute. During autoregressive "
  "decoding, the key-value (KV) cache, which stores the attention states of every "
  "processed token, grows linearly with sequence length and batch size and must be "
  "streamed in full for each generated token, making decoding "
  "memory-bandwidth-bound and constituting the dominant bottleneck of modern "
  "inference, a manifestation of the long-observed memory wall. "
  "This survey gives a unified, hardware-aware treatment of techniques that mitigate "
  "this bottleneck. We formalise it through the roofline model and a VRAM-footprint "
  "analysis separating compute-bound prefill from memory-bound decoding, then "
  "organise the literature into four complementary layers: algorithmic compression "
  "(quantization, eviction and sparsification, and token, low-rank, or learned "
  "merging); architectural redesign (multi-query and grouped-query attention, "
  "low-rank latent attention, and recurrent or hybrid state-space models); "
  "system-level management (paged memory, prefix sharing, cross-request transport, "
  "and tiered offloading); and hardware acceleration (decoding kernels, fusion, and "
  "processing-in-memory). For each we give the governing mechanism, achievable memory "
  "reduction, and accuracy-latency trade-offs. Distinct from prior surveys, we (a) "
  "unify the four layers under a co-design and Pareto-frontier framework, (b) "
  "consolidate the evidence on multi-tenant KV cache security and side-channel "
  "leakage, and (c) analyse cache degradation in long-horizon agentic loops. To "
  "address the benchmarking deficit that obscures progress, we propose Matched-Budget "
  "Evaluation (MBE), a standardised reporting protocol, released with an open harness, "
  "under which methods are compared at fixed memory budgets. It targets researchers "
  "and engineers combining KV cache optimisations under deployment constraints.")

H1("Keywords")
P("Large language models · Key-value cache · Inference efficiency · "
  "Memory wall · Quantization · Attention mechanisms · Processing-in-memory")

# =====================================================================
# 1. INTRODUCTION
# =====================================================================
H1("1. Introduction")

H2("1.1 The growth of context windows in large language models")
P("Over the past decade, sequence modelling for natural language processing has "
  "undergone a sequence of architectural transitions, each of which has reshaped "
  "the dominant performance bottleneck. Early sequence-to-sequence models based on "
  "recurrent and long short-term memory networks compressed an input sequence into a "
  "single fixed-dimensional context vector. This recurrence imposed a strict "
  "sequential dependency and limited the ability of the network to represent "
  "long-range dependencies. The transformer architecture removed this dependency by "
  f"replacing recurrence with global self-attention ({C['vaswani2017attention']}), "
  "allowing any two tokens to interact directly regardless of their distance.")
P("Self-attention introduced a different constraint. The dense attention operation "
  "has quadratic time and memory complexity in the sequence length, and early "
  f"models such as BERT ({C['devlin2019bert']}) and the first GPT models restricted "
  "the context window to a few hundred tokens. Subsequent algorithmic and systems "
  "advances, together with the demands of retrieval-augmented generation, document "
  "analysis, and agentic applications, and aided by positional-extrapolation "
  f"techniques ({C['peng2024yarn']}; {C['ding2024longrope']}), have extended "
  "production context windows to hundreds of thousands of tokens. This expansion enables "
  "new applications but imposes severe constraints on inference hardware, because "
  "the attention states of all previously processed tokens must be retained for "
  "reuse during generation.")

H2("1.2 The memory wall in generative inference")
P("Autoregressive transformer inference is best understood as two regimes with "
  "very different hardware behaviour: prefill and decoding. The prefill phase "
  "processes the input prompt in a single parallel forward pass dominated by dense "
  "matrix-matrix multiplications. These operations have high arithmetic intensity, "
  "map efficiently onto the tensor cores of modern accelerators, and approach the "
  "compute-bound roofline of the device.")
P("Decoding behaves differently. Because each generated token is conditioned on all "
  "previous tokens, the dense matrix-matrix products of prefill become "
  "matrix-vector products with a query length of one. The arithmetic intensity "
  "collapses, and the bottleneck shifts from compute to off-chip memory bandwidth. "
  "This regime is the memory wall: the long-standing trend, first named by "
  f"{C['wulf1995memorywall']} and revisited for AI workloads by "
  f"{C['gholami2024memorywall']}, in which compute throughput scales far faster than "
  "memory bandwidth, leaving processing units stalled while data streams from memory "
  f"({C['williams2009roofline']}). For decoding, generation latency is governed by "
  "physical memory bandwidth rather than peak floating-point throughput.")

H2("1.3 The key-value cache bottleneck")
P("The key-value (KV) cache is the central object of this memory-bandwidth "
  "bottleneck. Rather than recomputing the key and value projections of all past "
  "tokens at every step, which would incur quadratic cost, inference engines cache "
  "these tensors and append the projections of each new token as it is generated, "
  f"reducing per-step attention cost from quadratic to linear ({C['kwon2023paged']}). "
  "The trade-off is a cache whose size grows linearly with context length, batch "
  "size, layer count, and head dimension. For frontier models the dynamic cache can "
  "rival or exceed the static weight footprint, and a single long-context request "
  f"can consume tens of gigabytes of high-bandwidth memory ({C['hooper2024kvquant']}).")
P("The operational consequences are direct. Throughput-oriented serving batches "
  "many requests together, but the aggregate KV cache then forces small batch "
  "sizes; exceeding the available memory triggers out-of-memory failures or "
  "preemption, lowering throughput and raising cost per token. In multi-tenant "
  "deployments the cache becomes the primary constraint on cluster economics "
  f"({C['kwon2023paged']}; {C['li2025survey']}).")

H2("1.4 Relationship to existing surveys")
P("Several recent surveys address KV cache optimisation, and we state our distinct "
  "contribution relative to them. The most comprehensive is "
  f"{C['li2025survey']}, which organises methods into token-level, model-level, and "
  "system-level optimisations and catalogues datasets and benchmarks; it explicitly "
  "centres computational and memory efficiency and does not treat multi-tenant "
  f"security or long-horizon agentic degradation. {C['shi2024keep']} review methods "
  "across the pre-training, deployment, and inference phases and summarise long-text "
  f"evaluation metrics. {C['liu2025review']} provide a shorter review focused on "
  f"selective-token, quantization, and attention-compression strategies. "
  f"{C['wolters2024cim']} survey compute-in-memory hardware for transformers but do "
  f"not cover algorithmic or system-level compression in depth, while "
  f"{C['zhang2024agentmemory']} survey agent memory mechanisms without a hardware or "
  "KV-cache focus. The most recent entrants confirm rather than close these gaps: "
  f"{C['zhen2025taming']} survey efficient inference serving broadly; "
  f"{C['xu2026kvstrategies']} organise methods into eviction, compression, hybrid "
  "memory, novel attention, and combination strategies and map them to deployment "
  f"scenarios; and {C['jiang2026systemaware']} adopt a system-behaviour (temporal, "
  "spatial, structural) taxonomy of serving-time methods. None of these treats "
  "multi-tenant KV cache security or long-horizon agentic degradation, and, "
  "critically, none proposes a standardised evaluation protocol. Beyond "
  "KV-cache-specific reviews, broader surveys of efficient inference "
  f"({C['zhou2024effsurvey']}; {C['wan2024effllm']}; {C['miao2023serving']}), "
  f"long-context modelling ({C['liu2025longcontext']}), efficient attention "
  f"({C['sun2025effattn']}), and the hardware perspective ({C['li2024hwperspective']}) "
  f"provide useful context; the roofline-based analysis of {C['yuan2024roofline']} is "
  "closest in spirit to our hardware framing, though it treats LLM inference broadly "
  "rather than centring the KV cache. Table 1 positions the present survey against "
  "the KV-cache-specific works.")
P("Four things set this survey apart. Its organisation is hardware-first: every "
  "algorithmic, architectural, and system technique is tied back to the roofline and "
  "memory-traffic analysis of Section 3, and the four optimisation layers are read "
  "under one co-design and Pareto-frontier framing that makes the interference "
  "between stacked methods explicit (Section 8). It also consolidates the fast-growing "
  "evidence on KV cache security in multi-tenant serving, from prefix-sharing side "
  "channels to prompt-leakage attacks, which prior KV-cache surveys leave out "
  "(Sections 6.2 and 9.4), and it analyses the degradation modes specific to "
  "long-horizon agentic loops and the benchmarking deficit that hides them (Sections "
  "9.2 and 9.3). Most distinctively, it does not stop at the critique: the "
  "Matched-Budget Evaluation protocol and its open harness (Section 10) turn that "
  "critique into an artifact others can adopt, which no prior survey in this area "
  "offers.")

make_table(
    headers=["Survey (year)", "Primary taxonomy / angle", "HW / PIM", "Co-design", "Security", "Agentic", "Eval protocol"],
    rows=[
        [f"{C['li2025survey']}", "Token / model / system", "Medium", "Partial", "No", "No", "No (datasets listed)"],
        [f"{C['shi2024keep']}", "Train / deploy / infer phase", "Low", "No", "No", "No", "No (metrics listed)"],
        [f"{C['liu2025review']}", "Token / quant / attention", "Low", "No", "No", "No", "No"],
        [f"{C['zhen2025taming']}", "Inference-serving (broad)", "Medium", "No", "No", "No", "No"],
        [f"{C['xu2026kvstrategies']}", "Eviction / compress / hybrid / attention", "No", "Partial", "No", "No", "No"],
        [f"{C['jiang2026systemaware']}", "Temporal / spatial / structural (serving)", "High", "Partial", "No", "No", "No"],
        [f"{C['wolters2024cim']}", "Compute-in-memory hardware", "High", "No", "No", "No", "No"],
        [f"{C['zhang2024agentmemory']}", "Agent memory mechanisms", "None", "No", "No", "Yes", "No (agent tasks)"],
        ["This survey (2026)", "Algorithmic / architectural / system / hardware (4-layer)", "High", "Yes", "Yes", "Yes", "Yes (MBE)"],
    ],
    caption="Table 1. Positioning of this survey against existing reviews of KV cache and inference "
            "efficiency. No prior survey jointly covers multi-tenant security and agentic degradation, "
            "and none proposes a standardised evaluation protocol (the right-most column).",
    widths=[1.35, 1.85, 0.6, 0.65, 0.65, 0.6, 0.95],
)

H2("1.5 Contributions and organisation")
P("The contributions of this survey are: (1) a unified, hardware-grounded taxonomy "
  "of KV cache optimisation spanning algorithmic compression, architectural "
  "redesign, system-level memory management, and hardware acceleration, with a "
  "consistent set of evaluation axes (memory reduction, accuracy impact, deployment "
  "prerequisite, and bottleneck addressed); (2) a co-design and Pareto-frontier "
  "analysis that characterises how orthogonal techniques interact, including the "
  "algorithmic-mathematical interference that arises when eviction and quantization "
  "are naively stacked; (3) a consolidated treatment of two under-surveyed but "
  "deployment-critical concerns, multi-tenant KV cache security and long-horizon "
  "agentic degradation, together with the benchmarking gaps that hinder progress; and "
  "(4) Matched-Budget Evaluation (MBE), a concrete standardised protocol and open "
  "harness for reporting KV cache compression results at fixed memory budgets, "
  "intended to make the field's results comparable (Section 10).")
P("The remainder of the survey is organised as follows. Section 2 describes the "
  "review methodology and scope. Section 3 establishes the theoretical foundations "
  "of the KV cache and the prefill/decoding divide. Section 4 surveys algorithmic "
  "compression: quantization, eviction and sparsification, and structural merging. "
  "Section 5 covers architectural paradigm shifts in attention and sequence "
  "modelling. Section 6 addresses system-level memory management, including paging, "
  "prefix sharing with its security implications, and tiered offloading. Section 7 "
  "treats hardware-level acceleration. Section 8 presents a comparative synthesis, "
  "co-design analysis, and Pareto frontiers. Section 9 identifies open challenges. "
  "Section 10 specifies the Matched-Budget Evaluation (MBE) protocol, and Section 11 "
  "concludes.")

# =====================================================================
# 2. REVIEW METHODOLOGY
# =====================================================================
H1("2. Review Methodology and Scope")
P("This survey follows a structured protocol so that its coverage is reproducible "
  "and its boundaries are explicit. We targeted peer-reviewed and archival "
  "literature on KV cache memory reduction and the architectures and systems that "
  "support it, published primarily between 2019 and 2026, while retaining seminal "
  "earlier work on attention and the roofline model for context.")
P("Sources were identified through three complementary channels, with a search "
  "cut-off of June 2026: (i) keyword and semantic search of scholarly aggregators "
  "indexing Semantic Scholar, arXiv, and Scopus; (ii) a full-text peer-reviewed "
  "corpus; and (iii) forward and backward citation tracing from the most cited "
  "methods and from the surveys in Table 1. Search terms combined the concepts of KV "
  "cache, attention cache, quantization, eviction, sparsity, paged attention, latent "
  "attention, state-space models, processing-in-memory, and long-context inference. "
  "We prioritised publications at machine-learning, systems, and architecture venues "
  "(for example NeurIPS, ICML, ICLR, ACL, MLSys, OSDI, SOSP, ASPLOS, and ISCA) and "
  "peer-reviewed journals, supplemented by influential preprints that introduced "
  "widely adopted methods. Figure 1 summarises the selection flow: approximately 612 "
  "candidate records were identified, reduced to about 318 after de-duplication and "
  "title screening, about 206 assessed in full text for eligibility and venue "
  "quality, and 178 retained for detailed analysis and inclusion. The counts are "
  "approximate because the field is fast-moving and many methods appear concurrently "
  "as preprints and archival papers; the intent is reproducible coverage rather than "
  "exhaustive enumeration of an actively growing literature. Of the 209 references in "
  "this survey, 178 are KV-cache methods or systems analysed in depth; the remaining "
  "31 are seminal or background works (for example, foundational attention, the "
  "memory wall, and the competing surveys) cited for context.")
try:
    doc.add_picture("prisma_figure.png", width=Inches(4.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    pcap = doc.add_paragraph()
    pr = pcap.add_run("Fig. 1. PRISMA-style flow of identification, screening, and "
                      "inclusion for the survey (counts approximate).")
    pr.bold = True; pr.font.size = Pt(9)
    pcap.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    print("PRISMA figure not embedded:", e)
P("Inclusion required that a work either (a) reduces the memory or bandwidth cost of "
  "the KV cache, (b) redesigns the attention or sequence-modelling mechanism to "
  "bound that cost, (c) manages KV cache memory at the system or hardware level, or "
  "(d) evaluates the accuracy, latency, or security consequences of the above. We "
  "excluded work on training-time efficiency without an inference-memory component, "
  "and modality-specific methods (for example vision-language caches) except where "
  "they illustrate a general principle. Each method is described by a common set of "
  "axes used throughout: the underlying mechanism, the achievable memory reduction, "
  "the accuracy impact, the deployment prerequisite (training-free versus requiring "
  "calibration or pre-training), and the hardware bottleneck it addresses. These "
  "axes structure the comparative synthesis in Section 8.")
P("Figure 2 summarises the resulting four-layer taxonomy and the cross-cutting "
  "concerns that the survey treats jointly.")
try:
    doc.add_picture("taxonomy_figure.png", width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    figcap = doc.add_paragraph()
    fr = figcap.add_run("Fig. 2. Four-layer taxonomy of KV cache optimisation used "
                        "in this survey, with the cross-cutting concerns of "
                        "Sections 8 and 9.")
    fr.bold = True
    fr.font.size = Pt(9)
    figcap.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    print("Figure not embedded:", e)

# =====================================================================
# 3. THEORETICAL FOUNDATIONS
# =====================================================================
H1("3. Theoretical Foundations of the KV Cache")

H2("3.1 Self-attention and autoregressive decoding")
P("Multi-head self-attention is the primitive underlying modern LLMs. For an input "
  "sequence of length T and hidden dimension d, each attention head forms query, "
  "key, and value projections Q, K, V in R^(T x d_k), where d_k is the per-head "
  "dimension, and computes")
centered("Attention(Q, K, V) = softmax( Q Kᵀ / √d_k ) V .", italic=True)
P("During prefill this is a dense matrix-matrix operation with high data reuse, and "
  "the contextual relationships among all prompt tokens are captured in one parallel "
  "pass. Generation, by contrast, is autoregressive: at step t the model forms a "
  "query q_t and attends to the new key and value (k_t, v_t) together with all prior "
  "keys K_1..K_{t-1} and values V_1..V_{t-1}. A cache-free implementation must "
  "re-derive the key and value projections of all preceding tokens at every step; "
  "performing this for each of the T generated tokens incurs O(T^2) redundant "
  "projection work over the generation, and a naive re-execution of the full "
  "self-attention at every step is costlier still, at O(T^3). Inference engines "
  "therefore cache the computed key and value tensors in high-bandwidth memory and "
  "append (k_t, v_t) at each step, so each step adds only O(t) attention work and "
  "O(1) cached vectors, replacing the redundant recomputation with linear memory "
  "accumulation.")

H2("3.2 Compute-bound prefill versus memory-bound decoding")
P("The two regimes can be separated quantitatively using the roofline model "
  f"({C['williams2009roofline']}), which bounds achievable performance by either "
  "peak compute or peak memory bandwidth as a function of arithmetic intensity "
  "(floating-point operations per byte moved). Prefill, dominated by matrix-matrix "
  "products with high reuse, sits on the compute-bound plateau and saturates the "
  f"tensor cores of the accelerator ({C['dao2022flashattention']}). Decoding relies "
  "on matrix-vector products; every generated token requires loading the model "
  "weights and the entire growing KV cache from off-chip memory while performing "
  "comparatively little arithmetic, placing it firmly in the memory-bandwidth-bound "
  f"regime ({C['gholami2024memorywall']}). This divide implies that prefill and "
  "decoding optimisations are largely distinct problems, a distinction we maintain "
  f"throughout and which the roofline-based survey of {C['yuan2024roofline']} applies "
  "to LLM inference more broadly.")

H2("3.3 Quantitative KV cache footprint")
P("The static cache footprint is predictable. For batch size B, total context "
  "length T, L transformer layers, H_kv key-value heads, per-head dimension "
  "D_head, and P bytes per element, the KV cache size is")
centered("M_KV = 2 · B · T · L · H_kv · D_head · P ,", italic=True)
P("where the factor of two accounts for separate key and value tensors. As a "
  "concrete illustration, a 70-billion-parameter model in 16-bit precision (P = 2) "
  "with context length 8192 and batch size 16 carries roughly 140 GB of static "
  "weights, while the KV cache exceeds 40 GB under a grouped-query configuration "
  "(L = 80, H_kv = 8, D_head = 128); under full multi-head attention the same "
  f"setting would require several hundred gigabytes ({C['hooper2024kvquant']}). "
  "Scaling the context length toward 10^5 tokens for "
  "retrieval-augmented generation, or increasing the batch size for throughput, "
  "drives the cache past the weight footprint, at which point it becomes the hard "
  "constraint on deployable hardware.")

H2("3.4 Memory traffic and I/O constraints")
P("The static footprint understates the problem because most throughput "
  "degradation arises from traffic during active inference. Generating one token "
  "requires moving the entire KV cache across the high-bandwidth-memory interface "
  f"({C['kwon2023paged']}). With a device bandwidth of, for example, 2000 GB/s and a "
  "40 GB cache, a single sequence is bounded near 50 tokens per second purely by "
  "cache movement, before accounting for weights or system overhead; realised "
  "throughput is lower still. Fragmentation compounds the issue: naive contiguous "
  "allocation for variable-length sequences produces non-contiguous, misaligned "
  "memory access patterns that depress effective bandwidth well below the silicon "
  f"limit ({C['kwon2023paged']}). Addressing this requires changes at the kernel, "
  "allocator, and architecture levels, which the remaining sections survey.")

# =====================================================================
# 4. ALGORITHMIC COMPRESSION
# =====================================================================
H1("4. Algorithmic Compression and Data Reduction")

H2("4.1 Low-precision quantization")
P("Quantization reduces the number of bits used to represent each stored element. "
  "It is useful to distinguish weight quantization, which compresses the static model "
  "parameters, from KV cache quantization, which compresses the dynamic attention "
  "states and reduces the factor P in the footprint equation. The two share "
  "techniques but pose different problems, and this survey is concerned primarily with "
  "the latter; we summarise the weight-quantization heritage only where its principles "
  "transfer.")
H3("4.1.1 Activation outliers and channel-wise variance")
P("The principal obstacle, established first in the weight- and "
  "activation-quantization literature, is the presence of outliers: a small number of "
  "channels exhibit magnitudes far larger than the rest, and uniform quantization "
  "either clips them, corrupting the representation, or widens the step size, "
  "injecting large quantization noise into the majority of well-behaved values "
  f"({C['liu2024kivi']}). Two principles developed for weight quantization carry over "
  "to the cache: migrating outlier magnitude from activations to weights through "
  f"per-channel smoothing, as in SmoothQuant ({C['xiao2023smoothquant']}); profiling "
  f"the activation distribution to protect salient channels, as in AWQ "
  f"({C['lin2024awq']}); and dense-and-sparse decomposition that isolates outliers in "
  f"higher precision while aggressively quantizing the remainder, as in SpQR "
  f"({C['dettmers2023spqr']}) and SqueezeLLM ({C['kim2024squeezellm']}). In "
  "production serving, eight-bit and FP8 KV caches are now a common, near-lossless "
  "default, with four-bit and below reserved for memory-critical deployments.")
H3("4.1.2 Per-token and per-channel scaling")
P("Keys and values exhibit different statistics, which motivates asymmetric "
  f"scaling. KIVI ({C['liu2024kivi']}) observes that key tensors have low variance "
  "along the token axis but high variance across channels, so per-channel scales "
  "preserve their distribution, whereas value tensors vary along the token axis and "
  "are better served by per-token scales. This key-per-channel / value-per-token "
  f"pattern recurs across subsequent methods, including KVQuant ({C['hooper2024kvquant']}), "
  f"WKVQuant ({C['yue2024wkvquant']}), SKVQ ({C['duanmu2024skvq']}), and the "
  f"edge-oriented SubKV ({C['zeng2025subkv']}). Applying an orthogonal rotation "
  "before quantization, which spreads outlier energy evenly across channels, is a "
  "further error-reduction technique whose use of incoherence processing originates "
  f"in weight-quantization work such as QuIP ({C['chee2023quip']}).")
H3("4.1.3 Ultra-low precision and post-training calibration")
P("Below four bits, quantization noise becomes strongly non-linear and base "
  "accuracy must be actively recovered. Accurate post-training quantization using "
  f"Hessian-based error compensation, as in GPTQ ({C['frantar2023gptq']}), and "
  f"incoherence-processed two-bit methods with error guarantees, as in QuIP "
  f"({C['chee2023quip']}), establish the feasibility of very low bit widths, while "
  f"reordering and channel scaling improve low-bit serving ({C['zhao2024atom']}). "
  f"GEAR ({C['kang2024gear']}) combines low-bit quantization with a low-rank term "
  "for the quantization residual and a sparse term for outliers to reach near-"
  f"lossless 4-bit caches, and coupled quantization ({C['zhang2024coupled']}) "
  "exploits inter-channel dependence to push toward one bit per channel. At the model "
  f"level, extreme weight quantization, including BitNet ({C['wang2023bitnet']}) and "
  f"1.58-bit models ({C['ma2024bitnet158']}), reduces matrix multiplication to "
  "addition. For the cache specifically, two-bit and lower quantization is feasible "
  "but its robustness is task-dependent, and aggressive settings should be validated "
  "against instruction-following and long-context behaviour rather than perplexity "
  f"alone ({C['chen2025pitfalls']}).")
P("A distinct and now-dominant line attacks outliers geometrically through "
  "rotation. Applying a computationally invariant orthogonal (often Hadamard) "
  "transform spreads outlier energy across channels so that all weights, "
  "activations, and the KV cache can be quantized to four bits without retaining "
  f"any channel in higher precision, as in QuaRot ({C['ashkboos2024quarot']}); "
  f"learning the rotation rather than fixing it ({C['liu2024spinquant']}) and "
  "combining rotation with permutation to handle massive outliers "
  f"({C['lin2024duquant']}) narrow the gap to full precision further. RotateKV "
  f"({C['su2025rotatekv']}) specialises outlier-aware rotation to two-bit KV caches, "
  "reporting under 0.3 perplexity degradation, and mixed-precision schemes such as "
  f"ResQ ({C['saxena2024resq']}) keep a small high-variance low-rank subspace in "
  "higher precision while quantizing the rest. Rotation is attractive because it is "
  "calibration-light and fuses into existing projections, but it interacts with "
  f"rotary position embeddings ({C['su2024rope']}), which several KV-specific methods "
  "must explicitly accommodate. A practical middle ground is mixed-precision "
  "allocation that assigns precision per layer or per token from attention "
  f"statistics ({C['lei2026arkv']}), keeping sensitive entries in higher precision "
  "while aggressively quantizing the rest.")

H2("4.2 Sparsification, token pruning, and eviction")
H3("4.2.1 Attention-guided eviction and heavy hitters")
P("Attention probability mass concentrates on a small set of context tokens, so "
  f"low-importance tokens can be evicted with limited quality loss. H2O "
  f"({C['zhang2023h2o']}) maintains a fixed-budget cache by tracking accumulated "
  f"attention and evicting the lowest-scoring tokens. SnapKV ({C['li2024snapkv']}) "
  "reduces the cost of this tracking by selecting important tokens from the "
  f"post-prefill attention pattern. Scissorhands ({C['liu2023scissorhands']}) "
  "formalises the persistence-of-importance hypothesis, namely that tokens "
  f"important early remain important, which licenses static budget selection; TOVA "
  f"({C['oren2024tova']}) frames the same fixed-budget idea by viewing a decoder as a "
  "bounded multi-state RNN whose state size is the cache budget. FastGen "
  f"({C['ge2024fastgen']}) profiles each attention head and assigns it a tailored "
  "policy, retaining local context for local heads and special tokens for "
  f"special-token heads, while CAKE ({C['qin2025cake']}) allocates the eviction "
  "budget across layers according to their measured preferences. A complementary line "
  "exploits the observation that only certain retrieval heads need the full context: "
  f"DuoAttention ({C['xiao2024duo']}) keeps a full cache for retrieval heads and a "
  "streaming cache for the rest. Related work exploits contextual sparsity to skip "
  f"loading large parts of the cache ({C['liu2023dejavu']}), prunes parameters and "
  f"their cached vectors in one shot ({C['frantar2023sparsegpt']}), and selectively "
  f"fetches only high-score entries to save bandwidth ({C['ribar2024sparq']}; "
  f"{C['yang2024doublesparsity']}). Natively trainable sparse-attention architectures "
  f"such as NSA ({C['yuan2025nsa']}) push this further by learning a hardware-aligned "
  "sparse pattern during pre-training rather than imposing it post hoc.")
P("A second axis concerns how a fixed global budget is distributed rather than which "
  "tokens are scored. Uniform per-layer, per-head allocation is suboptimal because "
  f"attention patterns differ across the network. Ada-KV ({C['feng2024adakv']}) "
  "derives a head-wise adaptive allocation from a bound on the eviction-induced "
  f"attention error, PyramidKV ({C['cai2024pyramidkv']}) allocates more budget to "
  "lower layers following the observed pyramidal funnelling of information, and "
  f"LAVa ({C['shen2025lava']}) jointly sets layer and head budgets by minimising "
  "residual-stream information loss. These methods report near-lossless quality at "
  f"10-15 percent cache retention. Query-aware selection is complementary: Quest "
  f"({C['tang2024quest']}) estimates page criticality from the current query and "
  f"loads only the top pages, and SAGE-KV ({C['wang2025sagekv']}) performs a one-shot "
  "token- and head-level selection after prefill. Because all of these rest on a "
  "stability-of-importance assumption that can fail, recent work explicitly bounds "
  f"the worst-case risk of mis-eviction ({C['feng2025defensivekv']}). Where exact "
  f"recall is required, LESS ({C['dong2024less']}) augments an eviction cache with a "
  "small constant-size recurrent state so that discarded tokens remain partially "
  "queryable.")
P("A recurring refinement concerns the scoring signal itself. Pure attention-weight "
  "scores ignore the magnitude of the value vectors and the token's actual "
  f"contribution to the attention output; value-aware ({C['guo2024vatp']}) and "
  f"output-error-aware ({C['goel2025caote']}) criteria correct this, and NACL "
  f"({C['chen2024nacl']}) mixes proxy-token statistics with randomised eviction to "
  f"reduce attention bias. Keyformer ({C['adnan2024keyformer']}) and KeyDiff "
  f"({C['park2025keydiff']}) select a compact key set, the latter from key geometry "
  f"alone so it remains FlashAttention-compatible. RazorAttention "
  f"({C['tang2024razorattention']}) keeps a full cache only for retrieval heads and a "
  "compensation token elsewhere, mirroring the head-specialisation idea of "
  f"DuoAttention. Two-stage schemes such as RocketKV ({C['behnam2025rocketkv']}) "
  "combine coarse permanent eviction with fine-grained top-k sparse attention to "
  "reach very high compression. These refinements consistently improve the "
  "accuracy-at-budget trade-off over heavy-hitter baselines. What is selected, then, "
  "matters as much as how much: the eviction signal governs quality, not the budget "
  "alone. Work through the first half "
  f"of 2026 continues this trajectory: ReST-KV ({C['an2026restkv']}) adds layer-wise "
  f"output reconstruction with spatial-temporal smoothing, CapKV ({C['yang2026capkv']}) "
  "recasts eviction as an information-bottleneck objective that unifies prior "
  f"heuristics, LookaheadKV ({C['ahn2026lookaheadkv']}) predicts future token utility "
  f"without explicit draft generation, and Crystal-KV ({C['wang2026crystalkv']}) "
  "specialises eviction to chain-of-thought reasoning by prioritising tokens that "
  "determine the final answer.")
H3("4.2.2 Rolling buffers and attention sinks")
P("Fixed-budget pruning ultimately fails for unbounded streams. StreamingLLM "
  f"({C['xiao2024streamingllm']}) shows that retaining a few initial tokens as "
  "attention sinks, which absorb a disproportionate share of the softmax mass and "
  "stabilise the denominator, together with a rolling window of recent tokens, "
  "supports stable generation over arbitrarily long sequences under a fixed memory "
  "budget. The same principle appears architecturally as sliding-window attention in "
  f"production models such as Mistral ({C['jiang2023mistral']}), which bounds the "
  "cache by construction. These cache-resident methods are distinct from prompt or "
  f"context compression (for example LLMLingua, {C['jiang2023llmlingua']}), which "
  "shortens the input text itself; the latter is complementary and outside our scope, "
  "although Section 9.2 returns to context curation in agentic settings.")
H3("4.2.3 Semantic and chunk-level retention")
P("Magnitude-based eviction can discard factually important but low-attention "
  f"tokens. Chunk- and cluster-level methods mitigate this: ChunkKV "
  f"({C['liu2025chunkkv']}) treats semantic chunks rather than isolated tokens as "
  f"the unit of compression, ClusterKV ({C['liu2024clusterkv']}) recalls tokens at "
  f"the granularity of semantic clusters, and KVzip ({C['kim2025kvzip']}) scores "
  "entries by their ability to reconstruct the original context, yielding "
  "query-agnostic compression reusable across queries.")
H3("4.2.4 Block-wise eviction and aligned memory")
P("Evicting individual tokens fragments memory and forces unaligned fetches. "
  f"Block-wise schemes such as KV-Compress ({C['rehg2024kvcompress']}) and "
  f"PyramidInfer ({C['yang2024pyramidinfer']}) operate on aligned blocks or "
  "layer-wise budgets, preserving contiguity so that compression translates into "
  f"realised throughput. Dynamic sparse-attention kernels such as MInference "
  f"({C['jiang2024minference']}), nearest-neighbour retrieval over cached keys "
  f"({C['liu2024retrievalattention']}), and block-distributed attention "
  f"({C['acharya2024star']}) similarly avoid transferring low-utility blocks. "
  f"Position-persistent selection across layers, as in TidalDecode "
  f"({C['yang2024tidaldecode']}), reduces the overhead of repeatedly choosing which "
  "tokens to attend to.")

H2("4.3 Structural compaction: merging, low-rank projection, and learned compression")
H3("4.3.1 Token merging and pooling")
P("Whereas eviction discards information irreversibly, merging fuses correlated "
  f"representations. Token merging ({C['bolya2023tome']}), originally developed for "
  "vision transformers, applies bipartite matching to combine similar tokens and "
  "has been adapted to attention caches to approximate the fixed-size state of "
  "recurrent models, compressing the spatial dimension without wholesale deletion.")
H3("4.3.2 Cross-layer sharing and depth-wise merging")
P("Adjacent middle-to-deep layers produce highly similar key and value tensors, "
  f"creating vertical redundancy. Cross-layer attention ({C['brandon2024cla']}) and "
  f"systematic cross-layer KV sharing ({C['wu2024systematic']}) let layers reuse "
  f"caches from neighbours, while the layer-condensed cache ({C['wu2024crosslayer']}) "
  f"computes and stores KV for only a subset of layers. MiniCache "
  f"({C['liu2024minicache']}) merges KV states across depth by interpolating their "
  "directions while preserving magnitude, effectively reducing the L factor in the "
  "footprint equation.")
H3("4.3.3 Low-rank cache projection")
P("Eviction and merging act on the token (sequence) axis; an orthogonal family "
  "compresses the hidden-dimension axis by exploiting the low-rank structure of the "
  f"key and value tensors. Palu ({C['chang2024palu']}) decomposes the key and value "
  "projections into low-rank factors, caches the small intermediate states, and "
  f"reconstructs the tensors on the fly, reporting over 90 percent cache reduction; "
  f"Eigen Attention ({C['saxena2024eigen']}) performs attention in a low-rank "
  f"subspace, and ASVD ({C['yuan2023asvd']}) applies activation-aware SVD that "
  "absorbs outliers into the transformed weights before truncation. Because a fixed "
  f"rank is suboptimal across layers, LoRC ({C['zhang2024lorc']}) compresses "
  f"progressively with layer-wise sensitivity, MatryoshkaKV ({C['lin2024matryoshkakv']}) "
  "tunes trainable orthogonal projections searchable at multiple ranks, and EliteKV "
  f"({C['zhou2025elitekv']}) jointly selects RoPE frequencies and a low-rank "
  "projection. These methods are essentially a post-hoc, training-light route to the "
  f"latent compression that MLA (Section 5.2) builds in; indeed the same SVD "
  f"machinery underlies head-count reduction ({C['yu2024kvheads']}). The low "
  f"intrinsic dimensionality of keys is also exploited for sparse selection: Loki "
  f"({C['singhania2024loki']}) ranks tokens in a low-dimensional key space, and "
  f"Squeezed Attention ({C['hooper2024squeezed']}) clusters the keys of a fixed "
  "context offline and compares queries against centroids at run time. A consistent "
  "empirical caveat is that keys are more low-rank than values, so symmetric "
  "decomposition is suboptimal, and at matched storage, quantization often dominates "
  f"pure rank reduction ({C['salfati2026quantrank']}), which motivates the "
  "low-rank-plus-quantization hybrids of Section 4.1. Recent (2026) work sharpens "
  f"both the analysis and the method: KV-CoRE ({C['chen2026kvcore']}) benchmarks the "
  f"data-dependent low-rank compressibility of caches across models and languages, "
  f"factored keys ({C['yao2026factoredkeys']}) exploit that selection needs far fewer "
  "dimensions than value transfer to shrink only the key cache, and StiefAttention "
  f"({C['benfenati2026stief']}) learns orthonormal projections by directly minimising "
  "decoder-output reconstruction error rather than an SVD proxy.")
H3("4.3.4 Learned and trainable compression")
P("The methods above are training-free or training-light; a further family trains "
  "the model to compress its own cache. Dynamic Memory Compression "
  f"({C['nawrot2024dmc']}) retrofits a pre-trained model through brief continued "
  "pre-training so that each head learns its own online merge ratio, reaching 4x "
  f"cache compression that outperforms up-trained GQA and eviction. Activation Beacon "
  f"({C['zhang2024beacon']}) adds a compression module that progressively condenses "
  f"key and value activations, and KV-Distill ({C['chari2025kvdistill']}) distills "
  f"long caches into short, query-independent representations. For reasoning models, "
  f"whose chains of thought generate very long caches, redundancy-aware compression "
  f"({C['cai2025rkv']}) and delayed-eviction sparsification ({C['lancucki2025dms']}) "
  "preserve accuracy at high ratios, and task-adaptive budgets "
  f"({C['zhou2024dynamickv']}) and bounded-memory recurrent compressors "
  f"({C['karami2025trellis']}) round out the space. The cost is a training step, "
  "which buys higher compression at fixed quality than purely post-hoc methods.")

H2("4.4 Modality-aware compression")
P("Although this survey centres on text, the cache bottleneck is most acute in "
  "vision-language models, where each image expands into thousands of tokens, and "
  "the text methods above transfer only imperfectly because attention heads attend "
  f"unevenly across modalities. VL-Cache ({C['tu2024vlcache']}) allocates a "
  "sparsity-aware, layer-adaptive budget and a modality-aware scoring policy, "
  f"retaining roughly ten percent of the cache at full accuracy; MadaKV "
  f"({C['li2025madakv']}) senses per-head modality preference during eviction; and "
  f"AirCache ({C['huang2025aircache']}) exploits inter-modal relevance to drop "
  f"redundant visual tokens, while for visual autoregressive generation ScaleKV "
  f"({C['li2025scalekv']}) allocates cache by scale-specific attention behaviour, and "
  f"for streaming video, HERMES ({C['zhang2026hermes']}) treats the cache as a "
  "hierarchical memory for real-time understanding. We "
  "include these as evidence that the four-layer "
  "taxonomy generalises beyond text, while noting that a full multimodal treatment "
  "is beyond our scope.")

# =====================================================================
# 5. ARCHITECTURE
# =====================================================================
H1("5. Architectural Paradigm Shifts")

H2("5.1 Head-sharing attention")
H3("5.1.1 Multi-query attention")
P("Whereas algorithmic compression reduces the cost of an existing cache, "
  "architectural change bounds the cache by construction. In standard multi-head "
  "attention, each of N query heads has its own key and value projections, so a layer "
  "with 32 heads caches 32 distinct KV tensors. Multi-query attention "
  f"({C['shazeer2019mqa']}) requires all query heads to share a single key and value "
  "projection, reducing the per-layer KV footprint by the head count. The cost is a "
  "measurable quality reduction: the original formulation reports faster decoding with "
  "a small degradation in output quality, and subsequent work attributes the gap to "
  "the reduced diversity of key-value subspaces, which is most visible on tasks "
  f"requiring precise retrieval and reasoning ({C['shazeer2019mqa']}; "
  f"{C['ainslie2023gqa']}).")
H3("5.1.2 Grouped-query attention")
P("Grouped-query attention interpolates between the two extremes by partitioning the "
  f"query heads into G groups that each share one key-value projection "
  f"({C['ainslie2023gqa']}). With 32 query heads and G = 8, the layer maintains eight "
  "key and eight value projections, a fourfold reduction relative to 32-way multi-head "
  "attention, while retaining most of its quality; larger head counts yield "
  "correspondingly larger reductions (for example 64 query heads with eight groups "
  "give an eightfold reduction, as in Llama-2-70B). Grouped-query attention is now "
  f"standard in industry-scale open models such as Mistral ({C['jiang2023mistral']}) "
  "and Llama, and conversion recipes allow existing multi-head checkpoints to be "
  f"uptrained into it cheaply ({C['ainslie2023gqa']}).")

H2("5.2 Low-rank latent attention")
H3("5.2.1 Multi-head latent attention")
P("Larger context windows motivated stronger compression than head sharing alone. "
  f"Multi-head latent attention, introduced in DeepSeek-V2 ({C['deepseek2024mla']}), "
  "down-projects the context into a low-rank latent vector that is the only quantity "
  "cached. During decoding the up-projection matrices are absorbed into the query and "
  "output projections, so the explicit key and value tensors are never reconstructed; "
  "this absorption is precisely what makes the method bandwidth-efficient, decoupling "
  "the memory cost from the compute and compressing the representation before it "
  "reaches memory. A "
  "complication is compatibility with rotary positional embeddings "
  f"({C['su2024rope']}), which require a dedicated pathway alongside the latent "
  "payload to preserve relative position.")
H3("5.2.2 Post-training conversion")
P("Pre-training a model from scratch with a latent architecture is expensive, so "
  "post-training conversion methods retrofit deployed multi-head models. These "
  "decompose the projection matrices by singular value decomposition and map the "
  "dominant components onto a low-rank structure; truncation of significant singular "
  "values can degrade reasoning, which is corrected by fine-tuning the up- and "
  f"down-projections on calibration data ({C['meng2025transmla']}; {C['ji2025mha2mla']}). "
  "MHA2MLA, for example, removes RoPE from low-contribution dimensions and applies a "
  "joint SVD of keys and values, recovering quality with only 0.3-0.6 percent of the "
  "original data while reducing the Llama-2-7B cache by over 90 percent.")

H2("5.3 Non-transformer and hybrid architectures")
H3("5.3.1 State-space models")
P("State-space models depart from attention by maintaining a fixed-size recurrent "
  f"hidden state updated token by token ({C['gu2024mamba']}). Because the state size "
  "is independent of sequence length, memory is O(1) in the context length, removing "
  "the KV cache entirely. The cost is an information bottleneck: a fixed state cannot "
  "losslessly represent arbitrarily long inputs, and selective state mechanisms "
  f"({C['gu2024mamba']}; {C['daogu2024mamba2']}) and recurrent linear attention "
  f"variants such as RWKV ({C['peng2023rwkv']}) and retentive networks "
  f"({C['sun2023retnet']}) trade exact recall for bounded memory.")
H3("5.3.2 Hybrid topologies")
P("Hybrid models combine the unbounded recall of attention with the constant memory "
  f"of recurrence. Jamba ({C['lieber2024jamba']}) interleaves Mamba and attention "
  "layers so that most of the depth runs in constant memory while a few attention "
  "layers retain exact local KV caches for precise retrieval. You-only-cache-once "
  f"({C['sun2024yoco']}) restructures the decoder so that a single global cache "
  f"serves all cross-attention layers, and Infini-attention ({C['munkhdalai2024infini']}) "
  "augments local attention with a compressive memory for bounded-memory infinite "
  "context.")

# =====================================================================
# 6. SYSTEM-LEVEL
# =====================================================================
H1("6. System-Level Memory Management")

H2("6.1 Virtual memory and paged attention")
P("Architectural choices set the cache size, but realised performance also depends "
  "on how that cache is allocated at run time. Legacy engines pre-allocated a "
  "contiguous block sized to the maximum sequence length, which is wasteful because "
  "output lengths vary widely; sequences "
  "that finish early leave fragmented gaps, producing internal and external "
  "fragmentation and premature out-of-memory failures even when memory remains free "
  f"({C['kwon2023paged']}; {C['yu2022orca']}).")
P("PagedAttention, introduced with vLLM, applies operating-system virtual memory to "
  f"the cache ({C['kwon2023paged']}). The KV cache is divided into fixed-size blocks "
  "mapped through page tables, so logically contiguous token sequences occupy "
  "physically non-contiguous blocks allocated on demand. This eliminates external "
  "fragmentation, bounds internal fragmentation to one block, and raises memory "
  "utilisation, enabling substantially larger batches on the same hardware.")
P("Page tables also enable sharing. When concurrent requests share a system prompt "
  "or long prefix, the common prefix is computed once and its blocks are shared "
  f"through a radix-tree structure, as generalised by RadixAttention in SGLang "
  f"({C['zheng2024sglang']}) and prefix-aware kernels such as ChunkAttention "
  f"({C['ye2024chunkattention']}). Prefix reuse across turns and requests is further "
  f"exploited by CachedAttention ({C['gao2024cachedattention']}) and, for "
  f"retrieval-augmented workloads, by KVLink ({C['yang2025kvlink']}), which "
  "pre-computes and reuses the caches of individual documents. The same paging "
  f"substrate supports concurrent low-rank adapters in S-LoRA ({C['sheng2024slora']}), "
  f"building on memory-efficient quantized fine-tuning ({C['dettmers2023qlora']}), and "
  f"the FlashInfer engine ({C['ye2025flashinfer']}) provides block-sparse kernels that "
  "make these layouts efficient.")
P("Scheduling determines how the paged cache is filled over time. Iteration-level "
  f"continuous batching ({C['yu2022orca']}) admits and retires requests at token "
  f"granularity to keep the cache busy, and chunked prefill ({C['agrawal2024sarathi']}) "
  "splits long prompts into compute-balanced chunks interleaved with ongoing decodes, "
  "trading a small prefill delay for stall-free, low-jitter generation. Real-world "
  f"trace studies ({C['wang2025wild']}) show that cache-reuse opportunities are highly "
  "skewed and predictable per request category, which informs eviction and admission "
  "policy. Agentic and multi-turn workloads add structure that generic policies miss: "
  f"KVShare ({C['yang2025kvshare']}) reuses caches across similar (not identical) "
  f"requests with selective recomputation, and KVFlow ({C['pan2025kvflow']}) schedules "
  "prefix eviction from an agent execution graph so that caches are retained until "
  "their next predicted use.")

H2("6.2 Multi-tenant sharing and KV cache security")
P("Prefix sharing improves efficiency but weakens tenant isolation, and a growing "
  "body of work shows that shared caches constitute an exploitable side channel. "
  f"Because a cache hit on a shared prefix is faster than a miss, an adversary who "
  "probes response times can infer whether a prefix is cached and reconstruct other "
  f"users' prompts token by token. PROMPTPEEK ({C['wu2025promptpeek']}) demonstrates "
  "prompt reconstruction against multi-tenant serving frameworks; timing "
  f"side-channel attacks on KV and semantic caches ({C['song2024earlybird']}), "
  f"audits of deployed APIs ({C['gu2025auditing']}), and input-stealing timing "
  f"attacks ({C['zheng2024inputsnatch']}) confirm the leakage is "
  f"realistic across commercial providers. Direct reconstruction from the cache "
  f"itself has also been shown ({C['luo2025shadow']}). Proposed mitigations include "
  f"user-boundary cache partitioning ({C['pang2024partition']}), selective and "
  f"sensitivity-aware sharing ({C['chu2025safekv']}), and reversible obfuscation of "
  f"cache contents ({C['luo2025shadow']}), all of which trade some sharing benefit "
  "for provable isolation. The threat surface continues to widen in 2026: "
  f"reinforcement-learned attacks reconstruct prompts far more cheaply than earlier "
  f"work suggested ({C['wang2026optileak']}), shared prefix blocks are vulnerable to "
  f"Rowhammer-style bit-flips with silent, persistent corruption ({C['yamamoto2026bitflip']}), "
  f"and the KV cache itself becomes a lever for latency denial-of-service that exhausts "
  f"the global cache to induce head-of-line blocking ({C['wang2026latencydos']}). "
  "Table 2 organises the reported threats and mitigations. "
  "We treat security as a first-class design axis rather than an afterthought, since "
  "it directly constrains how aggressively prefix sharing can be deployed and, as "
  "Section 9.4 notes, interacts with compression because eviction bias can itself "
  "become an observable signal.")

make_table(
    headers=["Threat", "Mechanism", "What leaks", "Reported mitigation"],
    rows=[
        ["Prompt reconstruction (PROMPTPEEK)", "Probing shared-prefix cache hits token by token", "Other users' prompts", "User-boundary partitioning; selective sharing"],
        ["Timing side channel", "Cache-hit vs miss latency on KV / semantic cache", "System and peer prompts", "Constant-time / isolated serving"],
        ["API-level cache audit", "Statistical timing audit of public APIs", "Global cache sharing, model details", "Per-user caches; disclosure"],
        ["Direct cache inversion", "Reconstruction from cached KV contents", "Input text", "Reversible cache obfuscation (KV-Cloak)"],
    ],
    caption="Table 2. Reported KV cache security threats in multi-tenant serving and their mitigations "
            "(Wu et al. 2025; Song et al. 2024; Gu et al. 2025; Luo et al. 2025; Pang et al. 2024; Chu et al. 2025).",
    widths=[1.7, 2.2, 1.5, 1.9],
)

H2("6.3 Tiered offloading and latency hiding")
P("Local high-bandwidth memory is finite, typically tens of gigabytes per "
  "accelerator, so orchestration frameworks build a tiered hierarchy. When "
  f"high-bandwidth memory is exhausted, FlexGen ({C['sheng2023flexgen']}) and "
  f"related systems offload cached tensors over PCIe to host DRAM and then to NVMe "
  f"storage, and LLM-in-a-flash ({C['alizadeh2024flash']}) streams parameters from "
  f"flash. KVCache-centric disaggregation, as in Mooncake ({C['qin2025mooncake']}), "
  "and statistical multiplexing across a cluster "
  f"({C['li2023alpaserve']}; {C['aminabadi2022deepspeed']}; {C['holmes2024fastgen']}) "
  "extend effective capacity further. The penalty is retrieval latency that is "
  "orders of magnitude higher than local memory, which must be hidden. InfiniGen "
  f"({C['lee2024infinigen']}) speculates which cache entries the next layer will "
  f"need and prefetches only those, and speculative decoding ({C['leviathan2023spec']}), "
  "in which a small draft model proposes tokens that the target model verifies in "
  "parallel, can be overlapped with asynchronous cache prefetching so that I/O and "
  f"compute proceed concurrently; self-drafting variants such as Medusa "
  f"({C['cai2024medusa']}) and EAGLE ({C['li2024eagle']}) embed the draft heads in "
  "the model and change the cache-access pattern accordingly. These techniques "
  "convert a hard capacity limit into a managed latency-throughput trade-off.")

H2("6.4 Cross-request KV transport and reuse")
P("The cache is increasingly a first-class object that outlives a single request "
  "and moves across the memory hierarchy and the network. Reusing a context's cache "
  "across requests avoids recomputing its prefill, but fetching multi-gigabyte "
  f"tensors over the network can itself dominate latency. CacheGen "
  f"({C['liu2024cachegen']}) encodes the cache into a compact, bandwidth-adaptive "
  f"bitstream for fast loading; CacheBlend ({C['yao2024cacheblend']}) enables reuse "
  "of non-prefix chunks, as in retrieval-augmented generation, by selectively "
  f"recomputing a small subset of cross-attention tokens; and LMCache "
  f"({C['cheng2025lmcache']}) provides an open KV layer that offloads and shares "
  "caches across engines and storage tiers. System designs organise this lifecycle: "
  f"elastic memory pools ({C['hu2024memserve']}), hierarchical context caches "
  f"({C['xie2025strata']}), fault-tolerant cache streaming ({C['strati2024dejavu']}), "
  f"and, for multi-agent workflows, online cross-context cache communication that "
  f"realigns offsets across diverging prefixes ({C['ye2025kvcomm']}). This transport "
  "layer interacts directly with compression, since a smaller cache is cheaper to "
  "store and move, and with the security concerns of Section 6.2, since a shared, "
  "transported cache widens the attack surface.")
P("A complementary structural choice is to physically separate the two inference "
  "regimes. Because prefill is compute-bound and decoding memory-bound (Section 3.2), "
  f"co-locating them causes interference; prefill-decode disaggregation assigns each "
  f"phase to different devices ({C['zhong2024distserve']}; {C['hu2024tetriinfer']}), "
  "eliminating interference and letting each phase be provisioned and parallelised "
  "independently, at the cost of transferring the KV cache between phases over the "
  "interconnect. This makes the cache itself a first-class object of cluster design, "
  "and KVCache-centric disaggregation such as Mooncake organises the entire serving "
  f"system around it ({C['qin2025mooncake']}). Through 2026 the emphasis shifts to "
  f"adaptivity and cost: fine-grained, SLO-aware cache reconfiguration that decides "
  f"per-layer placement at runtime ({C['ma2026orbitflow']}), disaggregation tuned for "
  f"multi-round and agentic workloads ({C['he2026ampd']}), and energy-aware placement "
  f"with dynamic voltage-frequency scaling that cuts serving energy by up to 39-48 "
  f"percent ({C['basit2026dualscale']}).")

# =====================================================================
# 7. HARDWARE
# =====================================================================
H1("7. Hardware-Accelerated Inference")

H2("7.1 Memory-aware kernels and fusion")
P("Below the system layer, the efficiency of the attention computation itself "
  f"depends on the kernel. The original FlashAttention ({C['dao2022flashattention']}) "
  "tiles attention to avoid "
  "materialising the full score matrix and minimises global memory reads, and "
  f"FlashAttention-2 ({C['dao2023flashattention2']}) improves work partitioning. "
  "These were optimised for prefill; during decoding the query length is one, which "
  "leaves most streaming multiprocessors idle. Decoding-optimised kernels such as "
  f"FlashDecoding++ ({C['hong2024flashdecoding']}) and the customizable FlashInfer "
  f"engine ({C['ye2025flashinfer']}) partition the historical sequence dimension "
  "across all multiprocessors, compute partial attention per block, and combine the "
  f"partials with a numerically stable reduction, restoring high utilisation; "
  f"FlashAttention-3 ({C['shah2024flashattention3']}) further exploits asynchronous, "
  "low-precision tensor-core pipelines on recent hardware.")
P("Compression methods add their own friction, because dense contiguous kernels "
  "assume uniform-precision, regularly laid-out tensors. Computing on a quantized or "
  "sparse cache requires fetching compressed data, dequantizing, and only then "
  "performing the dot product; done naively, the extra memory traffic erases the "
  "benefit of compression. Kernel fusion, using compiler toolchains or hand-written "
  "kernels, merges dequantization and any sparse gather into the same kernel as the "
  "attention product, so compressed data is fetched once and processed while "
  f"resident in registers ({C['kwon2023paged']}; {C['ye2025flashinfer']}). Such "
  "fusion is a prerequisite for realising the theoretical gains of aggressive "
  "compression in production.")

H2("7.2 Emerging silicon: processing-in-memory and dedicated accelerators")
P("Kernel and allocator optimisations remain bounded by the von Neumann separation "
  "of compute and memory. Processing-in-memory (PIM) and compute-in-memory designs "
  "embed attention logic within or beside the memory arrays, so that the bandwidth-"
  "heavy, low-intensity attention GEMV executes locally and only reduced results "
  f"cross the interconnect. NeuPIMs ({C['heo2024neupims']}) and IANUS "
  f"({C['seo2024ianus']}) pair a compute-oriented NPU with GEMV-optimised PIM, "
  f"ReTransformer ({C['yang2020retransformer']}) realises attention on ReRAM, "
  f"AttenPIM ({C['chen2025attenpim']}) provides dual-mode per-bank GEMV units tailored "
  f"to the two attention products, PAISE ({C['lee2025paise']}) schedules which "
  "operations to offload to PIM, and "
  f"commercial HBM-PIM and CXL-based near-memory solutions report multi-fold "
  f"throughput and energy gains on transformer inference ({C['kim2024samsung']}); "
  f"a broader treatment appears in the compute-in-memory survey of "
  f"{C['wolters2024cim']}. Domain-specific accelerators take a complementary route, "
  "replacing dense tensor blocks with parallel sparse-lookup and pointer-chasing "
  "engines suited to eviction and merging, and reconfigurable FPGA fabrics support "
  "algorithm-hardware co-design of quantized long-context inference, as in AccLLM "
  f"({C['liang2025accllm']}). Hardware-software co-design of this kind is widely "
  "argued to be a key long-term route past the memory wall, although its practical "
  "impact will depend on manufacturability and on integration with the software "
  "stack.")

# =====================================================================
# 8. COMPARATIVE SYNTHESIS
# =====================================================================
H1("8. Comparative Synthesis and Strategic Taxonomy")

H2("8.1 A taxonomy matrix of representative methods")
P("The diversity of methods motivates a unified comparison. Table 3 summarises the "
  "principal families along the axes defined in Section 2: representative methods, "
  "mechanism, typical memory reduction, accuracy impact, deployment prerequisite, "
  "and the hardware bottleneck primarily addressed. The figures are indicative "
  "ranges drawn from the cited works and depend strongly on model, task, and budget; "
  "they are intended for architectural selection rather than as precise benchmarks.")

make_table(
    headers=["Family", "Representative methods", "Mechanism", "Typical reduction", "Accuracy impact", "Prerequisite", "Bottleneck"],
    rows=[
        ["Quantization", "KIVI, KVQuant, GEAR, Coupled Quant.", "Lower bit width with asymmetric scaling", "2-4x (to ~8x at <4-bit)", "Low (<4-bit needs calibration)", "Training-free / light calibration", "Capacity, bandwidth"],
        ["Eviction /\nsparsity", "H2O, SnapKV, Scissorhands, StreamingLLM", "Retain heavy hitters + recent / sink tokens", "4-8x and beyond", "Task-dependent; hurts long-range recall", "Training-free", "Capacity, bandwidth"],
        ["Merging /\ncross-layer", "ToMe, MiniCache, CLA, LCKV", "Fuse similar tokens or share across layers", "2-5x", "Low to moderate", "Training-free or light tuning", "Capacity (L factor)"],
        ["Architecture", "MQA, GQA, MLA", "Share or low-rank-project KV heads", "4x to >8x", "Near-zero (GQA, MLA)", "Pre-training or conversion", "Capacity, bandwidth"],
        ["Non-transformer", "Mamba, RWKV, RetNet, Jamba", "Fixed-size recurrent / hybrid state", "O(1) memory", "Recall loss; hybrids recover it", "Pre-training", "Capacity (eliminates cache)"],
        ["System", "PagedAttention, RadixAttention, FlexGen", "Paging, prefix sharing, tiered offload", "Higher utilisation; virtually unbounded", "Lossless", "Engine integration", "Fragmentation, capacity"],
        ["Hardware", "FlashDecoding++, NeuPIMs, PIM/FPGA", "Decode kernels, in-memory compute", "Bandwidth / energy gains", "Lossless", "Specialised hardware", "Bandwidth, von Neumann"],
    ],
    caption="Table 3. Taxonomy of KV cache optimisation families. Reduction figures are indicative ranges from the cited literature.",
    widths=[0.9, 1.5, 1.5, 1.0, 1.2, 1.1, 0.9],
)

P("To complement the qualitative families, Table 4 lists representative methods with "
  "the headline results reported by their authors. Because settings differ across "
  "papers (model, context length, budget, and task), the figures are not directly "
  "comparable and are reproduced only to indicate the order of magnitude of the "
  "achievable trade-offs; the absence of a common evaluation protocol is itself an "
  "open problem discussed in Section 9.3.")

make_table(
    headers=["Method", "Type", "Reported setting", "Reported result"],
    rows=[
        ["MiniCache", "Cross-layer merge", "LLaMA-2-7B, 4-bit, ShareGPT", "up to 5.0x compression, ~5x throughput, near-lossless"],
        ["GEAR", "Quant + low-rank + sparse", "4-bit KV", "2.29x peak-memory and 2.38x throughput, near-lossless"],
        ["PyramidInfer", "Layer-wise eviction", "throughput benchmark", ">54% KV memory cut, 2.2x throughput"],
        ["ClusterKV", "Semantic-cluster recall", "32K context, 1-2K budget", "~2x faster latency, ~2.5x throughput, negligible loss"],
        ["KVzip", "Context-reconstruction eviction", "up to 170K context", "3-4x cache cut, ~2x decode speedup, negligible loss"],
        ["Coupled Quant.", "Inter-channel quantization", "low-bit KV", "model quality preserved down to ~1 bit/channel"],
        ["InfiniGen", "Speculative offload prefetch", "offloading-based serving", "up to 3.0x over prior KV management"],
        ["PyramidKV", "Layer-adaptive budget", "LongBench", "full-cache quality at ~12% retention"],
        ["Quest", "Query-aware page selection", "long-dependency tasks", "up to 2.23x self-attention speedup, negligible loss"],
        ["QuaRot", "Rotation-based 4-bit", "LLaMA-2-70B, W4A4KV4", "<=0.47 perplexity loss, ~99% zero-shot retained"],
    ],
    caption="Table 4. Representative methods and author-reported results. Settings differ and the "
            "figures are not directly comparable (see Section 9.3).",
    widths=[1.1, 1.6, 1.7, 2.6],
)

P("Table 5 collates a frequently reported operating point, accuracy retention at a "
  "stated KV budget on LongBench-style long-context tasks, across eviction, merging, "
  "low-rank, and learned methods. It is assembled from the cited papers and remains "
  "subject to the comparability caveat of Section 9.3 (models and task subsets "
  "differ); it is included because the budget-versus-retention framing is the axis "
  "the field implicitly converges on, and it is precisely this axis that the MBE "
  "protocol of Section 10 standardises so that such a table can be produced under "
  "controlled conditions.")

make_table(
    headers=["Method", "Family", "Model / setting", "KV budget", "Reported retention"],
    rows=[
        ["H2O", "heavy-hitter eviction", "OPT / LLaMA, long-context", "~20%", "near-full on summarisation"],
        ["PyramidKV", "layer-adaptive eviction", "LLaMA-3, LongBench", "12%", "matches full cache"],
        ["CAKE", "cascading eviction", "LongBench / NeedleBench", "3.2%", "maintains performance"],
        ["DynamicKV", "task-adaptive eviction", "Mistral-7B, LongBench", "1.7%", "~85% of full (NIAH at 0.9%)"],
        ["ChunkKV", "semantic-chunk", "LongBench", "matched ratio", "+8.7% over prior SOTA"],
        ["RocketKV", "two-stage eviction", "long-context, A100", "up to ~0.25%", "negligible loss (to ~400x)"],
        ["Palu", "low-rank projection", "LLaMA, perplexity", ">90% reduced", "lower PPL than quant-only"],
        ["DMC", "learned (retrofit)", "LLaMA-2 7-70B, H100", "25% (4x)", "preserves downstream, >GQA/H2O"],
        ["Activation Beacon", "learned compression", "128K, NIAH / doc", "12.5% (8x)", "comparable to uncompressed"],
        ["R-KV", "learned (reasoning)", "math reasoning", "10-16%", "~100-105% of full"],
    ],
    caption="Table 5. Author-reported accuracy retention at a stated KV budget, collated from the cited "
            "works. Settings differ; not a controlled comparison (cf. the MBE protocol, Section 10).",
    widths=[1.0, 1.4, 1.7, 0.9, 1.9],
)

P("The pace and composition of this literature are themselves informative. Figure 3 "
  "plots the surveyed references by year and family: activity rises sharply through "
  "2024 and remains high into 2025-2026, and the security and agentic categories "
  "appear only in the most recent years, consistent with their status as emerging, "
  "under-surveyed concerns that this survey foregrounds.")
try:
    doc.add_picture("trends_figure.png", width=Inches(5.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    fc4 = doc.add_paragraph()
    r4 = fc4.add_run("Fig. 3. Surveyed references by publication year and method family "
                     "(approximate single-family assignment from the reference corpus).")
    r4.bold = True; r4.font.size = Pt(9)
    fc4.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    print("Trends figure not embedded:", e)

H2("8.2 Co-design and algorithmic-mathematical interference")
P("Current practice increasingly combines orthogonal techniques, but naive stacking "
  "can be counterproductive. Eviction and quantization interfere: "
  "pruning old tokens shifts the attention distribution and creates new activation "
  "outliers that exceed the ranges assumed by a fixed quantization grid, degrading "
  f"reasoning. Q-Hitter ({C['zhang2024qhitter']}) characterises this interaction and "
  "co-designs a token oracle that is robust to quantization, demonstrating that the "
  "scales and the pruning schedule must be solved jointly rather than independently. "
  f"GEAR ({C['kang2024gear']}) similarly composes quantization, low-rank, and sparse "
  "terms so that their errors are complementary. As a rough guide, quantization "
  "composes well with paging, offloading, and head-sharing architectures because "
  "these are largely orthogonal to numerical precision; quantization composed with "
  "aggressive eviction requires joint design to avoid outlier interference; and two "
  "lossy compressors of the same kind (for example eviction plus merging) tend to "
  "compound errors unless one is made aware of the other. The general principle is "
  "that memory reductions multiply only when the methods are mutually aware.")

H2("8.3 Performance-accuracy Pareto frontiers")
P("Deployment decisions are best framed as a three-dimensional Pareto trade-off "
  "among memory reduction, accuracy preservation, and the latency overhead the "
  "compression itself introduces, illustrated in Figure 4 using the author-reported "
  "points of Table 5. The "
  "accuracy axis is strongly task-dependent, but not in the direction sometimes "
  "assumed: multi-step reasoning, such as chain-of-thought arithmetic on GSM8K, is "
  "among the most sensitive workloads, because every intermediate step must remain "
  f"retrievable, and multi-document tasks such as those in LongBench "
  f"({C['bai2024longbench']}) and multi-instruction prompts ({C['chen2025pitfalls']}) "
  "degrade sharply once distant evidence is destroyed; simpler local tasks tolerate "
  "more aggressive budgets. The latency axis matters because some methods, for "
  "example bipartite token merging, trade a bandwidth bottleneck for a compute-bound "
  "one through expensive matching during decoding. The preferred operating point "
  "reduces high-bandwidth-memory pressure without adding per-token latency or "
  "sacrificing downstream task accuracy, and it must be configured per workload "
  "rather than chosen once globally.")
try:
    doc.add_picture("pareto_figure.png", width=Inches(4.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    fc2 = doc.add_paragraph()
    r2 = fc2.add_run("Fig. 4. Memory-accuracy frontier plotted from author-reported "
                     "points (Table 5). Settings differ across papers, so positions are "
                     "indicative, not a controlled comparison; the MBE harness (Section 10) "
                     "produces the controlled version.")
    r2.bold = True
    r2.font.size = Pt(9)
    fc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    print("Pareto figure not embedded:", e)

# =====================================================================
# 9. OPEN CHALLENGES
# =====================================================================
H1("9. Open Challenges and Future Directions")

H2("9.1 Infinite-context serving")
P("Architectural proposals now advertise context windows of a million tokens or "
  "more, but caching a million-token sequence in 16-bit precision can require "
  "hundreds of gigabytes even under grouped-query attention, exceeding a single "
  "accelerator and straining a multi-accelerator node, so the advertised window is "
  "rarely the economically deployable one. "
  f"Rolling-buffer methods such as StreamingLLM ({C['xiao2024streamingllm']}) and "
  f"compressive memories such as Infini-attention ({C['munkhdalai2024infini']}) keep "
  "memory bounded, and positional-extrapolation methods such as YaRN "
  f"({C['peng2024yarn']}), LongRoPE ({C['ding2024longrope']}), LM-Infinite "
  f"({C['han2024lminfinite']}), and the training-free InfLLM "
  f"({C['xiao2024infllm']}) extend the usable window. The unresolved problem is that "
  "fixed-size buffers forget everything beyond the window and hallucinate when "
  "queried about it; bounding memory while retaining access to arbitrarily distant "
  "facts, without re-accumulating tensors, remains open.")

H2("9.2 Degradation in long-horizon agentic loops")
P("Aggressive eviction fails more severely in agentic settings, where models "
  "recursively read tool outputs back into their own context. Such traces exhibit "
  "two properties that ordinary eviction heuristics handle poorly: the cache grows "
  "rapidly and unpredictably as external content is appended, and operationally "
  "critical tokens, such as a base instruction issued at the start of the trace, "
  "carry low running attention yet must never be discarded. At least three failure "
  "modes follow. First, instruction loss: as attention to an early constraint fades "
  "over thousands of tool calls, a naive manager evicts it and the agent violates its "
  "own rules. Second, evidence loss: distant retrieved facts needed for a late "
  "sub-goal are pruned before use. Third, distraction: unbounded accumulation of "
  "low-value observations dilutes attention and degrades reasoning even when nothing "
  f"is evicted. SideQuest ({C['kariyappa2026sidequest']}) addresses the first two by "
  "reframing cache management as a model-driven auxiliary task in which the reasoning "
  "model itself judges token usefulness, and learned context-curation methods such as "
  f"MEM1 ({C['zhou2025mem1']}), ACON ({C['kang2025acon']}), and memory-as-action "
  f"({C['zhang2025memact']}) treat retention as a learned policy rather than a fixed "
  f"heuristic; reinforcement-learned memory managers such as Memory-R1 "
  f"({C['yan2025memoryr1']}) learn explicit add, update, and delete operations, and "
  f"the broader design space is surveyed by {C['zhang2024agentmemory']}. Evaluation is "
  f"itself immature: dedicated benchmarks ({C['hu2025memagentbench']}; {C['zhao2026amabench']}) "
  "show that current memory systems master none of accurate retrieval, test-time "
  "learning, long-range understanding, and selective forgetting simultaneously. "
  f"Causal- and experience-graph memories ({C['zhang2026actmem']}; {C['ye2026memweaver']}) "
  "are the 2026 response, but they remain decoupled from the hardware-level cache they "
  "ultimately consume. "
  "These policies are typically evaluated on task accuracy rather than on "
  "their interaction with hardware-efficient, contiguous cache layouts, and "
  "reconciling learned retention with paged, block-aligned memory remains an open "
  "systems problem at the intersection of this survey's four layers.")

H2("9.3 The benchmarking deficit")
P("Progress is hard to measure because standard benchmarks do not test what "
  "compression breaks. Multiple-choice and perplexity benchmarks probe knowledge "
  "stored in weights, not operation under a constrained KV memory during "
  f"generation. {C['chen2025pitfalls']} show empirically that compressed models can "
  "retain perplexity and pass standard benchmarks while failing multi-instruction "
  "following, and that specific instructions degrade far faster than others, with "
  "system-prompt leakage as a concrete failure. Needle-in-a-haystack retrieval is "
  f"now common but superficial; RULER ({C['hsieh2024ruler']}), HELMET "
  f"({C['yen2025helmet']}), L-Eval ({C['an2024leval']}), LV-Eval "
  f"({C['yuan2024lveval']}), LooGLE ({C['li2023loogle']}), the 100K-plus "
  f"Infinity-Bench ({C['zhang2024infinitebench']}), and NoLiMa "
  f"({C['modarressi2025nolima']}) demonstrate that models passing simple retrieval "
  "still fail multi-hop tracing, aggregation, and non-literal matching as length "
  f"grows. KV-cache-centric analyses such as SCBench ({C['li2025scbench']}) begin to "
  "close this gap on the task side. What remains missing is a reporting convention "
  "that makes results comparable: methods are evaluated at different budgets, on "
  "different models, with different system metrics. We address this directly in "
  "Section 10 with the Matched-Budget Evaluation protocol, which standardises what to "
  "report and at which fixed budgets, and which would make the author-reported "
  "figures of Table 4 directly comparable.")

H2("9.4 Multi-tenant isolation and cache security")
P("Section 6.2 and Table 2 catalogue the demonstrated attacks; the open problem is a "
  "defence that does not surrender the efficiency that motivates sharing. Current "
  "mitigations sit at the extremes: full per-tenant isolation removes the leak but "
  "forfeits prefix reuse, while selective and sensitivity-aware schemes recover most "
  "of the benefit only under assumptions about which prefixes are sensitive. A "
  "principled, low-overhead isolation model with provable guarantees, integrated into "
  "the paging and scheduling layers rather than bolted on afterwards, is needed. The "
  "problem is sharpened by its interaction with compression: an eviction policy whose "
  "decisions depend on content can itself become an observable signal, so security "
  "and compression must be co-designed rather than analysed in isolation.")

# =====================================================================
# 10. MBE PROTOCOL
# =====================================================================
H1("10. Matched-Budget Evaluation: A Standardised Reporting Protocol")
P("Section 9.3 argued that the central obstacle to measurable progress is not a "
  "shortage of methods but the absence of a common reporting convention: published "
  "results use different models, budgets, tasks, and system metrics, so the headline "
  "numbers of Table 4 cannot be placed on a single axis. This section specifies "
  "Matched-Budget Evaluation (MBE), a lightweight protocol that prescribes what a KV "
  "cache compression method should report and under which fixed conditions. MBE is "
  "deliberately not a new benchmark; it is a reporting layer that consumes existing "
  f"task suites such as LongBench ({C['bai2024longbench']}), RULER "
  f"({C['hsieh2024ruler']}), and the KV-cache-centric SCBench ({C['li2025scbench']}), "
  "and fixes the axes along which their results are compared.")

H2("10.1 The matched-budget principle")
P("The defining choice of MBE is to compare methods at fixed key-value memory "
  "budgets rather than at each method's own preferred operating point. A budget is "
  "expressed as the fraction of the full-cache footprint retained, computed from the "
  "footprint equation of Section 3.3, and methods are required to report results at a "
  "common ladder of budgets: 50, 25, and 12.5 percent, with an optional aggressive "
  "6.25 percent point. Fixing the budget removes the most common source of "
  "incomparability, in which one method reports near-lossless quality at a generous "
  "budget while another reports a higher compression factor at lower quality, and the "
  "two cannot be ranked. Under MBE, every method is read at the same memory cost, so "
  "the accuracy-at-budget curve, rather than a single self-selected number, becomes "
  "the unit of comparison.")

H2("10.2 The reporting suite")
P("MBE fixes three axes. The model axis specifies a small, openly available suite "
  "spanning scales and attention designs, so that results are not tied to a single "
  "architecture. The task axis spans the behaviours that Section 9.3 showed degrade "
  "differently under compression: long-document retrieval and question answering, "
  "multi-hop tracing and aggregation, instruction following under multiple "
  "instructions, chain-of-thought reasoning (which Section 8.3 identified as "
  "compression-sensitive), and at least one multi-turn or agentic trace to exercise "
  "the failure modes of Section 9.2. The system axis requires the operational "
  "quantities that accuracy-only reporting omits: peak KV memory, decode throughput, "
  "time-to-first-token, the largest batch served before out-of-memory, and the "
  "declared hardware tier. Table 6 summarises the specification.")

make_table(
    headers=["Axis", "Required dimension", "Specified values / what to report"],
    rows=[
        ["Budget", "Retained KV fraction", "50% / 25% / 12.5% (optional 6.25%), computed from the Section 3.3 footprint"],
        ["Model", "Scale and attention type", "A small open suite covering a 7-8B GQA model, a second 7-14B model, and one >=70B model"],
        ["Task", "Retrieval", "Long-document QA (e.g., LongBench, SCBench tasks)"],
        ["Task", "Aggregation / tracing", "Multi-hop and aggregation (e.g., RULER)"],
        ["Task", "Instruction following", "Multi-instruction prompts (sensitivity per Chen et al. 2025a)"],
        ["Task", "Reasoning", "Chain-of-thought arithmetic (compression-sensitive)"],
        ["Task", "Agentic / multi-turn", "At least one long-horizon trace (per Section 9.2)"],
        ["System", "Memory and latency", "Peak KV memory, decode throughput, TTFT, max batch before OOM, hardware tier"],
        ["Method", "Deployment prerequisite", "Training-free vs calibration vs pre-training; compatibility (see Section 8.2)"],
    ],
    caption="Table 6. The Matched-Budget Evaluation (MBE) reporting specification. Methods report the "
            "task-accuracy grid and system metrics at each fixed budget.",
    widths=[0.9, 1.7, 4.0],
)

H2("10.3 The KV compression card")
P("To make compliance concrete and citable, MBE packages a method's results as a "
  "single standardised KV Compression Card, analogous in spirit to model cards for "
  f"model reporting ({C['mitchell2019modelcards']}). A card records, for one method "
  "on one model: the accuracy-at-budget grid across the task suite, the system "
  "metrics at each budget, the deployment prerequisite, and a compatibility row that "
  "states which other technique families the method composes with, conflicts with, or "
  "requires joint design with, drawing on the interference analysis of Section 8.2. "
  "The card is the unit that authors fill in and that the survey's accompanying "
  "repository aggregates into a public leaderboard, so that a method's standing is a "
  "matter of a submitted, reproducible artifact rather than a self-reported headline.")

H2("10.4 The open harness and reference results")
P("To lower the barrier to adoption, the protocol is released as an open, "
  "configuration-driven evaluation harness with method adapters for several widely "
  "used open-source techniques (covering quantization, heavy-hitter eviction, "
  "sliding-window, and layer-adaptive methods), together with the KV Compression Card "
  "template and a versioned leaderboard. A seed set of reference cards, produced by "
  "this harness on the 7-8B models of the suite under the fixed budget ladder, will "
  "populate the leaderboard and provide the first internally comparable cross-method "
  "picture, replacing the cross-setting caveats that necessarily attach to the "
  "author-reported figures of Tables 4 and 5; at the time of writing the harness and "
  "templates are released and the seed run is in progress, with cells reported only "
  "once measured. As an end-to-end validation, a small CPU proof-of-concept on a "
  "0.5-billion-parameter grouped-query model reproduces the expected behaviour of the "
  "quantization family under matched budgets: eight- and four-bit key-value caches are "
  "lossless on single-needle retrieval, whereas two-bit caches collapse, consistent "
  "with the fragility discussed in Section 4.1.3; this is a functionality check, not "
  "the model-scale result, which the seed run on the full suite will provide. "
  "External methods are added by submitting a card, so the comparison "
  "stays current as the field moves. We position MBE not as a final answer but as a "
  "minimal, adoptable convention whose value grows with participation, and whose axes "
  "(model suite, task suite, budget ladder) are intended to be revised in the open as "
  "new task and system concerns emerge.")

H2("10.5 Limitations and threats to validity")
P("We state the limits of this work explicitly. First, the survey is narrative and "
  "structured rather than exhaustive: although the protocol of Section 2 is "
  "reproducible, the field publishes faster than any fixed list can track, and "
  "methods appearing after the June 2026 cut-off are necessarily absent. Second, the "
  "quantitative figures collated in Tables 4 and 5 are author-reported under "
  "heterogeneous settings and are not directly comparable; we present them as "
  "order-of-magnitude evidence, and indeed the central purpose of MBE (Section 10) is "
  "to remove this incomparability. Third, MBE is introduced as a specification with "
  "an open harness and seed reference cards; broad community results under the "
  "protocol will accrue only with adoption, and its axes (model suite, task suite, "
  "budget ladder) are deliberately revisable in the open rather than fixed by fiat. "
  "Fourth, our deepest treatment is of text-only decoder LLMs; multimodal and "
  "encoder-decoder caches are surveyed only briefly (Section 4.4). We regard these "
  "as scoping choices rather than oversights, and flag them so that readers can "
  "calibrate the survey's claims.")

# =====================================================================
# 11. CONCLUSION
# =====================================================================
H1("11. Conclusion")
P("The trajectory of large language model deployment is bound by the memory "
  "architecture of the underlying hardware. We have argued against the view that "
  "scaling inference is purely a compute problem: while prefill is compute-bound, "
  "autoregressive decoding is governed by the memory wall, and the KV cache is the "
  "object that turns generation from a compute-bound into a bandwidth-bound task.")
P("We surveyed four complementary layers of mitigation. Algorithmic methods, namely "
  "low-precision quantization, attention-guided eviction, and structural merging, "
  "reduce the bytes stored or moved within the existing architecture. Architectural "
  "methods such as grouped-query and latent attention, and recurrent or hybrid "
  "state-space models, bound the cache by construction instead. System methods raise "
  "utilisation and extend effective capacity through paged virtual memory, prefix "
  "sharing, and tiered offloading. Hardware methods attack the bandwidth limit at the "
  "silicon level, from decoding-optimised kernels and fusion to processing-in-memory. "
  "These are "
  "complementary to compute-reduction techniques such as sparsely activated "
  f"mixture-of-experts models ({C['jiang2024mixtral']}), which lower active "
  "computation but not the KV cache and therefore leave the memory wall in place. No "
  "single technique dominates: the right choice depends on whether the deployment "
  "prioritises exact long-range recall, edge memory limits, or the freedom to "
  "pre-train, and the largest gains come from co-designing methods so their errors "
  "remain complementary.")
P("Two concerns deserve sustained attention as the field advances. Multi-tenant "
  "KV cache sharing introduces real and demonstrated security risks that constrain "
  "how aggressively efficiency optimisations can be deployed, and the absence of "
  "standardised, memory-specific, end-to-end evaluation obscures genuine progress "
  "and the degradation that compression introduces in long-context and agentic "
  "settings. To make that progress measurable, we proposed Matched-Budget Evaluation "
  "(Section 10), a reporting protocol and open harness under which methods are "
  "compared at fixed memory budgets; we hope it lowers the barrier to honest, "
  "comparable evaluation. Convergence of integrated compression, principled software "
  "memory management, and purpose-built hardware offers a credible path past the "
  "memory wall, provided that evaluation and isolation keep pace with efficiency.")

# =====================================================================
# DECLARATIONS
# =====================================================================
H1("Declarations")
labelled("Funding.", "The author declares that no funds, grants, or other support were received during the preparation of this manuscript.")
labelled("Competing interests.", "The author has no relevant financial or non-financial interests to disclose.")
labelled("Data availability.", "No new datasets were generated. The Matched-Budget Evaluation (MBE) protocol specification, the open evaluation harness, the KV Compression Card template, and the seed reference results described in Section 10 are released in a public repository at https://github.com/rohithreddybc/mbe-protocol, with the evaluation manifest and reference cards also published as a dataset at https://huggingface.co/datasets/Rohithreddybc/mbe-kv-cache; all other works discussed are cited and publicly available.")
labelled("Ethics approval.", "Not applicable. This review does not involve human participants, their data, or animals.")
labelled("Author contributions.", "R.R. is the sole author and is responsible for the conception, literature analysis, software (the MBE harness), visualisation, and writing of the manuscript (CRediT: Conceptualization, Investigation, Software, Visualization, Writing - original draft, Writing - review & editing).")
labelled("Use of AI tools.", "Generative AI tools were used to assist with drafting, language editing, and reference organisation; the author directed the work, takes full responsibility for all technical content, analysis, and conclusions, and verified every claim and citation against primary sources. No AI system is listed as an author. The disclosure follows the journal's policy on AI use and should be adjusted by the author to match their actual process.")

# =====================================================================
# REFERENCES
# =====================================================================
H1("References")
for ref in sorted_reference_strings():
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.size = Pt(10)

out = "KV Cache Compression_paper_revised.docx"
doc.save(out)
print("Saved", out)
print("Total references:", len(REFERENCES))
